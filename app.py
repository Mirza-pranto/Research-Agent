import json
import uuid
import httpx
import streamlit as st
import io
from docx import Document

st.set_page_config(page_title="AI Research Workspace", layout="wide")

if "sessions" not in st.session_state: st.session_state.sessions = {}
if "active_thread_id" not in st.session_state: st.session_state.active_thread_id = str(uuid.uuid4())

active_thread = st.session_state.active_thread_id
current_session_data = st.session_state.sessions.get(active_thread, {})

# Sidebar logic
with st.sidebar:
    st.title("📚 Research History")
    if st.button("➕ New Research Session", use_container_width=True):
        st.session_state.active_thread_id = str(uuid.uuid4())
        st.rerun()
    st.divider()
    if st.session_state.sessions:
        st.write("**Saved Sessions:**")
        for tid, sess in st.session_state.sessions.items():
            topic_label = sess.get("topic", "Untitled Session")
            btn_prefix = "🟢 " if tid == st.session_state.active_thread_id else "📄 "
            if st.button(f"{btn_prefix}{topic_label[:26]}", key=f"btn_{tid}", use_container_width=True):
                st.session_state.active_thread_id = tid
                st.rerun()

st.title("Free AI Research Agent Workspace 5.0")
st.write("Now with **Live Streaming & Document RAG!**")

default_prompt = current_session_data.get("topic", "")
prompt = st.text_area("Research prompt", value=default_prompt, placeholder="Enter a topic...", key=f"prompt_{active_thread}")
auto_approve = st.toggle("Auto-approve research plan", value=True)

# Helper function to generate Word Document
def generate_word_document(markdown_text: str) -> bytes:
    doc = Document()
    doc.add_heading('Research Report', 0)
    
    for line in markdown_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        
        # Map markdown to Word styles
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', '').strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', '').strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', '').strip(), level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        else:
            # Remove bold/italic markdown asterisks for cleaner Word text
            clean_line = line.replace('**', '').replace('*', '')
            doc.add_paragraph(clean_line)
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# Stream consumer function
def consume_stream(url: str, payload: dict, final_data: dict, status_box, streaming_placeholder=None):
    has_error = False
    accumulated_draft = ""
    
    try:
        with httpx.stream("POST", url, json=payload, timeout=600.0) as response:
            response.raise_for_status()

            for line in response.iter_lines():
                if not line or not line.startswith("data: "): continue
                raw_json = line[6:].strip()
                if not raw_json: continue

                event = json.loads(raw_json)
                event_type = event.get("type")

                # Handle Live Tokens
                if event_type == "token" and streaming_placeholder is not None:
                    accumulated_draft += event.get("content", "")
                    streaming_placeholder.markdown(accumulated_draft)
                    continue

                # Handle Node Updates
                if event_type == "update":
                    node = event.get("node")
                    new_status = event.get("status")
                    if new_status and new_status != "processing":
                        final_data["status"] = new_status

                    if node == "planner":
                        status_box.update(label="🧠 **Planner:** Research plan generated!", state="running")
                        final_data["plan"] = event.get("plan")
                    elif node == "retriever":
                        status_box.update(label="🌐 **Retriever:** Web sources collected!", state="running")
                        final_data["sources"] = event.get("sources", [])
                    elif node == "synthesizer":
                        status_box.update(label="✍️ **Synthesizer:** Research draft written!", state="running")
                        final_data["draft"] = event.get("draft")
                    elif node == "fact_checker":
                        status_box.update(label="✅ **Fact Checker:** Complete!", state="complete")
                    elif node == "error":
                        status_box.update(label="❌ Error occurred", state="error")
                        st.error(f"Agent Error: {event.get('message')}")
                        has_error = True
                    
    except Exception as err:
        status_box.update(label="❌ Connection failed", state="error")
        st.error(f"Backend Error: {err}")
        has_error = True
    
    # Ensure the final streamed draft is saved to state if it wasn't caught in the final update chunk
    if accumulated_draft and not final_data.get("draft"):
        final_data["draft"] = accumulated_draft

    return final_data, has_error

# Primary Run Button
if st.button("Run research", key=f"run_{active_thread}") and prompt.strip():
    status_box = st.status("🚀 Initializing research workflow...", expanded=True)
    streaming_placeholder = st.empty() # Create an empty container for live text

    final_data = {"topic": prompt.strip(), "plan": None, "sources": [], "draft": None, "status": "started"}
    payload = {"topic": prompt.strip(), "thread_id": active_thread, "auto_approve": auto_approve}

    final_data, has_error = consume_stream("http://localhost:8000/research/stream", payload, final_data, status_box, streaming_placeholder)
    st.session_state.sessions[active_thread] = final_data
    if not has_error: st.rerun()

# Interrupt UI
if current_session_data.get("status") == "planned":
    st.warning("✋ **Graph Paused:** Please review and edit the research questions.")
    plan_dict = current_session_data.get("plan", {})
    questions = plan_dict.get("questions", [])
    
    with st.form(key=f"approve_form_{active_thread}"):
        edited_qs_text = st.text_area("Target Search Questions", value="\n".join(questions), height=150)
        
        if st.form_submit_button("Approve & Resume Search", type="primary"):
            plan_dict["questions"] = [q.strip() for q in edited_qs_text.split("\n") if q.strip()]
            current_session_data["plan"] = plan_dict
            
            status_box = st.status("▶️ Resuming workflow...", expanded=True)
            streaming_placeholder = st.empty() # Container for live text upon resuming
            
            payload = {"thread_id": active_thread, "plan": plan_dict}
            current_session_data, has_error = consume_stream("http://localhost:8000/research/resume", payload, current_session_data, status_box, streaming_placeholder)
            st.session_state.sessions[active_thread] = current_session_data
            if not has_error: st.rerun()

# ==========================================
# COMPLETED STATE: Display Draft & Chat & Export
# ==========================================
if current_session_data and current_session_data.get("status") in ["completed", "verified", "drafted"]:
    st.divider()
    
    # 1. Display the Final Streamed Draft
    st.subheader("📝 Final Research Draft")
    draft_text = current_session_data.get("draft") or "No draft generated."
    st.markdown(draft_text)
    
    with st.expander("View Reference Sources"):
        for source in current_session_data.get("sources", []):
            st.markdown(f"- **{source.get('title', 'Unknown')}**: {source.get('url', '')}")

    st.divider()
    
    # ==========================================
    # UPGRADE: Deep-Dive Exporting UI
    # ==========================================
    st.subheader("📥 Export Report")
    col1, col2 = st.columns(2)
    
    with col1:
        st.download_button(
            label="📄 Download as Markdown",
            data=draft_text,
            file_name=f"Research_Report_{active_thread[:8]}.md",
            mime="text/markdown",
            use_container_width=True
        )
        
    with col2:
        word_file = generate_word_document(draft_text)
        st.download_button(
            label="📝 Download as Word (.docx)",
            data=word_file,
            file_name=f"Research_Report_{active_thread[:8]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    st.divider()

    # 2. UPGRADE: Chat with the Research
    st.subheader("💬 Chat with your Research")
    st.caption("Ask questions about the gathered sources.")
    
    # Initialize chat history for this thread
    chat_key = f"chat_{active_thread}"
    if chat_key not in st.session_state:
        st.session_state[chat_key] = []
        
    # Render previous messages
    for msg in st.session_state[chat_key]:
        st.chat_message(msg["role"]).write(msg["content"])
        
    # Chat Input
    if chat_prompt := st.chat_input("Ask a follow-up question..."):
        st.session_state[chat_key].append({"role": "user", "content": chat_prompt})
        st.chat_message("user").write(chat_prompt)
        
        with st.chat_message("assistant"):
            response_placeholder = st.empty()
            accumulated_response = ""
            
            # Connect to our new /research/chat endpoint
            payload = {"thread_id": active_thread, "message": chat_prompt}
            with httpx.stream("POST", "http://localhost:8000/research/chat", json=payload, timeout=60.0) as r:
                for line in r.iter_text():
                    accumulated_response += line
                    response_placeholder.write(accumulated_response)
            
            st.session_state[chat_key].append({"role": "assistant", "content": accumulated_response})